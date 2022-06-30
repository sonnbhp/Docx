from docx import Document
document = Document()
document.add_heading("BẢN KÊ KHAI",0)
document.add_heading("Kính gửi: Trưởng Phòng PC& KSNB",1)
p = document.add_paragraph("Họ và tên:")
p.add_run("Nguyễn Bá Sơn").bold = True
document.save("test.docx")
